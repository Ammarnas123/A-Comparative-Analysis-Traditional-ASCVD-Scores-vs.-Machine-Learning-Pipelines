import docx
from docx.shared import Pt
import os
from sklearn.model_selection import train_test_split, GridSearchCV
from imblearn.over_sampling import SMOTE
from sklearn.neighbors import KNeighborsClassifier
from sklearn.tree import DecisionTreeClassifier
from sklearn.metrics import confusion_matrix, accuracy_score, recall_score
import matplotlib.pyplot as plt
import numpy as np

def create_project_log():
    # Create a new Document
    doc = docx.Document()
    
    # Add Title
    title = doc.add_heading('Health Data Science Capstone Project Log', 0)
    title.alignment = 1 # Center align

    # Add Project Overview
    doc.add_heading('1. Project Overview & Objectives', level=1)
    doc.add_paragraph(
        "Objective: Create a machine learning tool that calculates the Atherosclerotic Cardiovascular "
        "Disease (ASCVD) risk predicting 10-year CVD events using the Framingham dataset (4,239 participants). "
        "The project compares the traditional ASCVD formula against modern Machine Learning algorithms."
    )
    doc.add_paragraph("Algorithms Planned: K-Nearest Neighbors (KNN), Decision Trees (DT), Random Forest, "
                      "Gradient Boosting, Logistic Regression, and Artificial Neural Networks (ANN).")
    doc.add_paragraph("Data Split: 70% Training / 30% Testing.")
    
    # Methodology
    doc.add_heading('2. Methodology', level=1)
    # Phase 1
    doc.add_heading('Phase 1: Exploratory Data Analysis (EDA) & Data Preprocessing', level=2)
    doc.add_paragraph(
        "- Data Exploration: Assess class imbalances, feature distributions, and correlations.\n"
        "- Missing Data Handling: Impute missing values using K-Nearest Neighbors (KNN Imputation) for robust estimation.\n"
        "- Scaling and Normalization: Apply standard scaling to continuous variables."
    )

    # Phase 2
    doc.add_heading('Phase 2: Traditional ASCVD Risk Calculation', level=2)
    doc.add_paragraph(
        "- Baseline Calculation: Use traditional cohort equations to calculate the 10-year risk.\n"
        "- Thresholding: Convert the continuous risk to a binary prediction (>7.5% or >20%).\n"
        "- Deviation Analysis: Calculate baseline error against actual `TenYearCHD` outcomes."
    )
    
    # Phase 3
    doc.add_heading('Phase 3: Machine Learning Model Development (70/30 Split)', level=2)
    doc.add_paragraph(
        "- Feature Engineering & Selection: Forward Feature Selection and Backward Feature Elimination.\n"
        "- Train Algorithms: KNN, Decision Trees, Random Forest, Gradient Boosting, etc."
    )
    
    # Phase 4
    doc.add_heading('Phase 4: Model Evaluation', level=2)
    doc.add_paragraph(
        "- Metrics: Confusion Matrix, Sensitivity (Recall), Specificity, F1-Score, ROC-AUC."
    )
    
    # Step-by-Step Execution Log
    doc.add_heading('3. Step-by-Step Execution Log', level=1)
    
    # Step 1: EDA Log
    doc.add_heading('Step 1: Exploratory Data Analysis (EDA)', level=2)
    doc.add_paragraph('Status: In Progress\nDate: [Insert Date Here]', style='Intense Quote')
    doc.add_paragraph(
        "Action: Loaded the Framingham dataset (4,239 rows) to inspect missing variables, "
        "target variable distribution, and correlation matrices."
    )
    
    # Add EDA Code Block
    doc.add_heading('Python Code Used for EDA:', level=3)
    code_paragraph = doc.add_paragraph(
        "import pandas as pd\n"
        "import numpy as np\n"
        "import matplotlib.pyplot as plt\n"
        "import seaborn as sns\n\n"
        "def perform_eda():\n"
        "    sns.set_theme(style=\"whitegrid\")\n"
        "    file_path = r\"C:\\Users\\HP\\Downloads\\AML- CapstoneProject\\Framingham_CVD assesment\\framingham.csv\"\n"
        "    df = pd.read_csv(file_path)\n"
        "    # Code inspects df.shape, df.isnull().sum(), correlations, and outputs visualizations.\n"
        "    # For full code, view the Framingham_CVD_assesment.py script.\n\n"
        "if __name__ == '__main__':\n"
        "    perform_eda()"
    )
    code_paragraph.style.font.name = 'Courier New' # Make it look like code
    
    # Placeholders for future steps
    doc.add_heading('Step 2: Data Preprocessing & Missing Value Imputation', level=2)
    doc.add_paragraph('Status: Pending', style='Intense Quote')
    
    doc.add_heading('Step 3: Calculating Baseline ASCVD Risk', level=2)
    doc.add_paragraph('Status: Pending', style='Intense Quote')
    
    doc.add_heading('Step 4: Machine Learning Modeling (KNN & Decision Trees)', level=2)
    doc.add_paragraph('Status: Completed', style='Intense Quote')
    doc.add_paragraph("Action: Split data 70/30 (Training/Testing). Applied SMOTE to balance the training data. "
                      "Used GridSearchCV to find the optimal hyperparameters for KNN (K) and Decision Trees (Alpha).")
    doc.add_paragraph("Results for K-Nearest Neighbors:", style='List Bullet')
    doc.add_paragraph("Best K found: 5", style='List Bullet 2')
    doc.add_paragraph("Accuracy: 64.94% | Recall: 48.19%", style='List Bullet 2')

    doc.add_paragraph("Results for Decision Tree:", style='List Bullet')
    doc.add_paragraph("Best Alpha (ccp_alpha) found: 0.0050", style='List Bullet 2')
    doc.add_paragraph("Accuracy: 50.86% | Recall: 75.65%", style='List Bullet 2')

    doc.add_paragraph("Interpretation: The Decision Tree pruned with Alpha=0.0050 sacrificed some accuracy but achieved a massive jump in Recall to 75.65%, making it much safer clinically than KNN. Interestingly, the optimized tree completely discarded most features, relying almost entirely on just two variables to make its predictions.")

    doc.add_paragraph("Top Clinical Variables (Feature Importance):", style='List Bullet')
    doc.add_paragraph("1. Age (80.03% importance)", style='List Bullet 2')
    doc.add_paragraph("2. Systolic Blood Pressure / sysBP (19.97% importance)", style='List Bullet 2')

    doc.add_heading('Step 5: Advanced Modeling (LogReg, Random Forest, Gradient Boosting)', level=2)
    doc.add_paragraph('Status: Completed', style='Intense Quote')
    doc.add_paragraph("Action: Trained Logistic Regression, Random Forest, and Gradient Boosting Classifiers on the 70/30 split. "
                      "Compared accuracy vs. recall trade-offs.")
    doc.add_paragraph("Results for Logistic Regression (Clinical Baseline):", style='List Bullet')
    doc.add_paragraph("Accuracy: 66.82% | Recall: 61.14%", style='List Bullet 2')
    doc.add_paragraph("Results for Random Forest:", style='List Bullet')
    doc.add_paragraph("Accuracy: 74.14% | Recall: 43.52%", style='List Bullet 2')
    doc.add_paragraph("Results for Gradient Boosting (GBM):", style='List Bullet')
    doc.add_paragraph("Accuracy: 79.56% | Recall: 18.13%", style='List Bullet 2')

    doc.add_paragraph("Interpretation: There is a massive clinical trade-off occurring. Gradient Boosting achieved the highest overall accuracy (almost 80%), "
                      "but its recall crashed to 18%, meaning it missed 82% of the sick patients. Logistic Regression provided the most balanced, safe approach "
                      "(66.8% Acc, 61.1% Recall). However, the absolute best model for purely catching sick patients remains the heavily pruned single Decision Tree from Step 4 (75% Recall).")

    doc.add_paragraph("Top Variables (Random Forest Importance):", style='List Bullet')
    doc.add_paragraph("1. Age (24.10%)", style='List Bullet 2')
    doc.add_paragraph("2. Systolic BP (14.60%)", style='List Bullet 2')
    doc.add_paragraph("3. Total Cholesterol (9.50%)", style='List Bullet 2')
    doc.add_paragraph("4. Diastolic BP (9.41%)", style='List Bullet 2')
    doc.add_paragraph("5. Glucose (8.50%)", style='List Bullet 2')

    doc.add_heading('Step 6: Feature Engineering, Scaling & Final Artificial Neural Network (ANN)', level=2)
    doc.add_paragraph('Status: Completed', style='Intense Quote')
    doc.add_paragraph("Action: Engineered 4 novel clinical features (Pulse Pressure, MAP, Age x Cigs, Age x sysBP). "
                      "Scaled the dataset using StandardScaler. Trained a Multi-Layer Perceptron (ANN), and re-trained the Decision Tree and Logistic Regression.")
    doc.add_paragraph("Results for Decision Tree (Engineered):", style='List Bullet')
    doc.add_paragraph("Accuracy: 70.83% | Recall: 57.51% (The new features significantly boosted the DT's accuracy from 50% to 70%, but sacrificed its industry-leading 75% recall).", style='List Bullet 2')
    doc.add_paragraph("Results for Logistic Regression (Engineered):", style='List Bullet')
    doc.add_paragraph("Accuracy: 67.30% | Recall: 60.62% (The model remained consistently balanced, solidifying its place as the safest, most stable predictor).", style='List Bullet 2')
    doc.add_paragraph("Results for Artificial Neural Network / MLP:", style='List Bullet')
    doc.add_paragraph("Accuracy: 73.35% | Recall: 41.97% (The ANN acted similarly to the Random Forest: high overall accuracy, but struggled significantly to identify the minority class of actual sick patients even with SMOTE data).", style='List Bullet 2')

    doc.add_heading('Conclusion & Presentation Generation', level=2)
    doc.add_paragraph('Status: Ready for Capstone PPT', style='Intense Quote')
    
    # Save Document
    save_path = r"C:\Users\HP\Downloads\AML- CapstoneProject\Framingham_CVD assesment\Capstone_Project_Log.docx"
    doc.save(save_path)
    print(f"✅ Word Document successfully created and saved at: {save_path}")

def train_baseline_models(df_clean):
    plt.close('all')
    print("\n=======================================================", flush=True)
    print("--- Phase 3: Optimized ML Baseline (KNN & Decision Tree) ---", flush=True)
    print("=======================================================\n", flush=True)
    
    # 1. Prepare Data
    X = df_clean.drop('TenYearCHD', axis=1)
    y = df_clean['TenYearCHD']
    
    # 70/30 Split
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.30, random_state=42, stratify=y)
    
    print("Applying SMOTE to balance the minority class in the training set...", flush=True)
    smote = SMOTE(random_state=42)
    X_train_smote, y_train_smote = smote.fit_resample(X_train, y_train)
    
    # --- MODEL 1: K-Nearest Neighbors (Optimizing K) ---
    print("\n--- Training Model 1: Optimizing K-Nearest Neighbors ---", flush=True)
    knn = KNeighborsClassifier()
    # Test these K values
    knn_params = {'n_neighbors': [3, 5, 7, 9, 11, 15, 21]}
    # We optimize for ROC AUC (distinguishing sick from healthy) instead of pure accuracy
    knn_grid = GridSearchCV(knn, knn_params, cv=5, scoring='roc_auc', n_jobs=-1)
    
    print("Searching for the best K...", flush=True)
    knn_grid.fit(X_train_smote, y_train_smote)
    best_knn = knn_grid.best_estimator_
    print(f"✅ Best KNN Model Found: K = {best_knn.n_neighbors}", flush=True)
    
    y_pred_knn = best_knn.predict(X_test)
    print(f"KNN Accuracy: {accuracy_score(y_test, y_pred_knn)*100:.2f}%", flush=True)
    print(f"KNN Recall:   {recall_score(y_test, y_pred_knn)*100:.2f}%\n", flush=True)
    
    # --- MODEL 2: Decision Tree (Optimizing Alpha) ---
    print("\n--- Training Model 2: Optimizing Decision Tree (ccp_alpha) ---", flush=True)
    dt = DecisionTreeClassifier(random_state=42, class_weight='balanced')
    # Test these Alpha values (0 means no pruning, higher means simpler trees)
    dt_params = {'ccp_alpha': [0.0, 0.001, 0.002, 0.005, 0.01, 0.02, 0.05]}
    dt_grid = GridSearchCV(dt, dt_params, cv=5, scoring='roc_auc', n_jobs=-1)
    
    print("Searching for the best Alpha...", flush=True)
    # Note: Decision Trees don't strictly need SMOTE as much if class_weights='balanced' is used, 
    # but we feed it the raw X_train so it learns true feature relationships natively.
    dt_grid.fit(X_train, y_train)
    best_dt = dt_grid.best_estimator_
    print(f"✅ Best Decision Tree Model Found: Alpha (ccp_alpha) = {best_dt.ccp_alpha:.4f}", flush=True)
    
    y_pred_dt = best_dt.predict(X_test)
    print(f"Decision Tree Accuracy: {accuracy_score(y_test, y_pred_dt)*100:.2f}%", flush=True)
    print(f"Decision Tree Recall:   {recall_score(y_test, y_pred_dt)*100:.2f}%\n", flush=True)
    
    # --- FEATURE IMPORTANCE EXTRACTION (FOR YOUR PPT) ---
    print("=======================================================", flush=True)
    print("--- 🚀 Top Clinical Variables (Feature Importance) 🚀 ---", flush=True)
    print("=======================================================", flush=True)
    
    # Extract importances from the best Decision Tree
    importances = best_dt.feature_importances_
    features = X.columns
    
    # Sort them descending
    indices = np.argsort(importances)[::-1]
    
    for i in range(len(features)):
        # Only print if it actually contributed > 0%
        if importances[indices[i]] > 0:
            print(f"{i+1}. {features[indices[i]]}: {importances[indices[i]]*100:.2f}% importance", flush=True)
    print("\n(Save these for your final PowerPoint!)\n", flush=True)
    print("=======================================================", flush=True)

if __name__ == "__main__":
    create_project_log()
